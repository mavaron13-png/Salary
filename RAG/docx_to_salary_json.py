from __future__ import annotations

import argparse
import json
import re
from collections import defaultdict
from pathlib import Path
from typing import Any

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P

SOURCE_NAME = "MyDNA Guía Salarial 2025"
COUNTRY = "Colombia"
CURRENCY = "COP"
UNIT = "Salario Bruto Mensual"

COMPANY_REVENUE_MAP = {
    "Empresa Pequeña": "<30MM USD",
    "Empresa Mediana": "30MM a 100MM USD",
    "Empresa Grande": ">100MM USD",
}


def iter_block_items(document: Document):
    """
    Itera párrafos y tablas en el orden real del documento.
    python-docx, en su infinita sabiduría burocrática, no lo hace fácil por defecto.
    """
    body = document.element.body

    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def normalize_text(value: str | None) -> str:
    if value is None:
        return ""

    return re.sub(r"\s+", " ", value).strip()


def clean_section_title(text: str) -> str:
    """
    Convierte:
    'Página 54 - Finanzas - Empresa Pequeña'
    en:
    'Finanzas - Empresa Pequeña'
    """
    text = normalize_text(text)
    text = re.sub(r"^Página\s+\d+\s*-\s*", "", text, flags=re.IGNORECASE)
    return text


def parse_section_title(text: str) -> tuple[str | None, str | None]:
    """
    Detecta títulos como:
    Finanzas - Empresa Pequeña
    Tecnología de la Información - Empresa Grande
    Ventas & Marketing - Empresa Mediana
    """
    text = clean_section_title(text)

    pattern = re.compile(
        r"^(?P<area>.+?)\s*-\s*(?P<size>Empresa\s+(?:Pequeña|Mediana|Grande))$",
        flags=re.IGNORECASE,
    )

    match = pattern.match(text)

    if not match:
        return None, None

    area = normalize_text(match.group("area"))
    size_raw = normalize_text(match.group("size")).lower()

    size_map = {
        "empresa pequeña": "Empresa Pequeña",
        "empresa mediana": "Empresa Mediana",
        "empresa grande": "Empresa Grande",
    }

    company_size = size_map.get(size_raw)

    return area, company_size


def parse_int(value: str | int | None) -> int | None:
    """
    Convierte:
    '$2,708,275' -> 2708275
    '2.708.275' -> 2708275
    '2708275' -> 2708275
    'N/A' -> None
    """
    if value is None:
        return None

    value = str(value).strip()

    if value.upper() in {"N/A", "NA", "", "-", "NULL", "NONE"}:
        return None

    digits = re.sub(r"[^\d-]", "", value)

    if digits in {"", "-"}:
        return None

    return int(digits)


def build_record(
        *,
        area: str,
        company_size: str,
        cargo: str,
        minimo: str,
        medio: str,
        maximo: str,
        variable: str,
        anual: str,
        source_file: str,
) -> dict[str, Any]:
    salario_minimo = parse_int(minimo)
    salario_medio = parse_int(medio)
    salario_maximo = parse_int(maximo)
    variable_int = parse_int(variable)
    salario_anual = parse_int(anual)

    facturacion = COMPANY_REVENUE_MAP.get(company_size)

    record_id = make_record_id(
        country=COUNTRY,
        area=area,
        company_size=company_size,
        cargo=cargo,
    )

    texto_rag = (
        f"En {COUNTRY}, según {SOURCE_NAME}, para el área de {area}, "
        f"en {company_size} con facturación anual {facturacion}, "
        f"el cargo {cargo} tiene salario bruto mensual mínimo de "
        f"{salario_minimo} {CURRENCY}, salario medio de {salario_medio} {CURRENCY}, "
        f"salario máximo de {salario_maximo} {CURRENCY}, variable {variable_int} "
        f"y salario anual de {salario_anual} {CURRENCY}."
    )

    return {
        "id": record_id,
        "pais": COUNTRY,
        "fuente": SOURCE_NAME,
        "archivo_fuente": source_file,
        "area": area,
        "tamano_empresa": company_size,
        "facturacion_anual": facturacion,
        "cargo": cargo,
        "salario_minimo_cop": salario_minimo,
        "salario_medio_cop": salario_medio,
        "salario_maximo_cop": salario_maximo,
        "variable": variable_int,
        "salario_anual_cop": salario_anual,
        "moneda": CURRENCY,
        "unidad": UNIT,
        "texto_rag": texto_rag,
    }


def make_record_id(
        *,
        country: str,
        area: str,
        company_size: str,
        cargo: str,
) -> str:
    raw = f"{country}-{area}-{company_size}-{cargo}".lower()

    raw = (
        raw.replace("á", "a")
        .replace("é", "e")
        .replace("í", "i")
        .replace("ó", "o")
        .replace("ú", "u")
        .replace("ñ", "n")
    )

    raw = re.sub(r"[^a-z0-9]+", "-", raw)
    raw = raw.strip("-")

    return raw


def parse_docx_tables(docx_path: Path) -> list[dict[str, Any]]:
    document = Document(docx_path)

    current_area = None
    current_company_size = None
    records: list[dict[str, Any]] = []

    for block in iter_block_items(document):
        if isinstance(block, Paragraph):
            text = normalize_text(block.text)

            area, company_size = parse_section_title(text)

            if area and company_size:
                current_area = area
                current_company_size = company_size

        elif isinstance(block, Table):
            if not current_area or not current_company_size:
                continue

            rows = block.rows

            if not rows:
                continue

            for row in rows[1:]:
                cells = [normalize_text(cell.text) for cell in row.cells]

                if len(cells) < 6:
                    continue

                cargo, minimo, medio, maximo, variable, anual = cells[:6]

                if not cargo or cargo.lower() == "cargo":
                    continue

                record = build_record(
                    area=current_area,
                    company_size=current_company_size,
                    cargo=cargo,
                    minimo=minimo,
                    medio=medio,
                    maximo=maximo,
                    variable=variable,
                    anual=anual,
                    source_file=docx_path.name,
                )

                records.append(record)

    return records


def is_value_token(token: str) -> bool:
    token = token.strip()

    if token.upper() in {"N/A", "NA"}:
        return True

    return bool(re.fullmatch(r"[$]?[0-9.,]+", token))


def parse_docx_plaintext_fallback(docx_path: Path) -> list[dict[str, Any]]:
    """
    Fallback por si el DOCX no trae tablas reales sino texto plano.
    Detecta líneas tipo:
    Analista Snr 2708275 3563519 4347494 0 49889271
    """
    document = Document(docx_path)

    lines = []

    for paragraph in document.paragraphs:
        text = normalize_text(paragraph.text)

        if text:
            lines.append(text)

    current_area = None
    current_company_size = None
    records: list[dict[str, Any]] = []

    for line in lines:
        area, company_size = parse_section_title(line)

        if area and company_size:
            current_area = area
            current_company_size = company_size
            continue

        if line.lower().startswith("cargo "):
            continue

        if not current_area or not current_company_size:
            continue

        parts = line.split()

        if len(parts) < 6:
            continue

        last_five = parts[-5:]

        if not all(is_value_token(token) for token in last_five):
            continue

        cargo = " ".join(parts[:-5])
        minimo, medio, maximo, variable, anual = last_five

        record = build_record(
            area=current_area,
            company_size=current_company_size,
            cargo=cargo,
            minimo=minimo,
            medio=medio,
            maximo=maximo,
            variable=variable,
            anual=anual,
            source_file=docx_path.name,
        )

        records.append(record)

    return records


def deduplicate_records(records: list[dict[str, Any]]) -> list[dict[str, Any]]:
    seen = {}

    for record in records:
        seen[record["id"]] = record

    return list(seen.values())


def build_table_chunks(records: list[dict[str, Any]]) -> list[dict[str, Any]]:
    grouped = defaultdict(list)

    for record in records:
        key = (record["area"], record["tamano_empresa"])
        grouped[key].append(record)

    table_chunks = []

    for (area, company_size), rows in grouped.items():
        facturacion = COMPANY_REVENUE_MAP.get(company_size)

        lines = [
            f"{COUNTRY} | {area} | {company_size} | Facturación anual: {facturacion}",
            "Cargo | Salario mínimo COP | Salario medio COP | Salario máximo COP | Variable | Salario anual COP",
        ]

        for row in rows:
            lines.append(
                f"{row['cargo']} | "
                f"{row['salario_minimo_cop']} | "
                f"{row['salario_medio_cop']} | "
                f"{row['salario_maximo_cop']} | "
                f"{row['variable']} | "
                f"{row['salario_anual_cop']}"
            )

        chunk_id = make_record_id(
            country=COUNTRY,
            area=area,
            company_size=company_size,
            cargo="tabla-completa",
        )

        table_chunks.append(
            {
                "id": chunk_id,
                "tipo_chunk": "tabla",
                "pais": COUNTRY,
                "fuente": SOURCE_NAME,
                "area": area,
                "tamano_empresa": company_size,
                "facturacion_anual": facturacion,
                "moneda": CURRENCY,
                "unidad": UNIT,
                "texto": "\n".join(lines),
                "record_ids": [row["id"] for row in rows],
            }
        )

    return table_chunks


def build_row_chunks(records: list[dict[str, Any]]) -> list[dict[str, Any]]:
    row_chunks = []

    for record in records:
        row_chunks.append(
            {
                "id": f"{record['id']}-chunk",
                "tipo_chunk": "fila",
                "pais": record["pais"],
                "fuente": record["fuente"],
                "area": record["area"],
                "tamano_empresa": record["tamano_empresa"],
                "cargo": record["cargo"],
                "moneda": record["moneda"],
                "unidad": record["unidad"],
                "texto": record["texto_rag"],
                "record_id": record["id"],
            }
        )

    return row_chunks


def save_json(path: Path, data: Any) -> None:
    path.write_text(
        json.dumps(data, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def save_jsonl(path: Path, rows: list[dict[str, Any]]) -> None:
    with path.open("w", encoding="utf-8") as file:
        for row in rows:
            file.write(json.dumps(row, ensure_ascii=False) + "\n")


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Convierte tablas salariales de DOCX a JSON maestro para RAG."
    )

    parser.add_argument(
        "--input",
        required=True,
        help="Ruta del archivo DOCX de entrada.",
    )

    parser.add_argument(
        "--output",
        default="salary_master.json",
        help="Ruta del JSON maestro de salida.",
    )

    parser.add_argument(
        "--records-jsonl",
        default="salary_records.jsonl",
        help="Ruta del JSONL con un registro salarial por línea.",
    )

    parser.add_argument(
        "--chunks-jsonl",
        default="salary_chunks.jsonl",
        help="Ruta del JSONL con chunks para RAG.",
    )

    args = parser.parse_args()

    docx_path = Path(args.input)

    if not docx_path.exists():
        raise FileNotFoundError(f"No existe el archivo: {docx_path}")

    records_from_tables = parse_docx_tables(docx_path)

    if records_from_tables:
        records = records_from_tables
    else:
        records = parse_docx_plaintext_fallback(docx_path)

    records = deduplicate_records(records)

    table_chunks = build_table_chunks(records)
    row_chunks = build_row_chunks(records)

    master = {
        "schema_version": "1.0",
        "fuente": SOURCE_NAME,
        "archivo_fuente": docx_path.name,
        "pais": COUNTRY,
        "moneda": CURRENCY,
        "unidad": UNIT,
        "total_registros": len(records),
        "total_table_chunks": len(table_chunks),
        "total_row_chunks": len(row_chunks),
        "records": records,
        "chunks": {
            "tables": table_chunks,
            "rows": row_chunks,
        },
    }

    save_json(Path(args.output), master)
    save_jsonl(Path(args.records_jsonl), records)
    save_jsonl(Path(args.chunks_jsonl), table_chunks + row_chunks)

    print(f"OK - registros extraídos: {len(records)}")
    print(f"JSON maestro: {args.output}")
    print(f"JSONL registros: {args.records_jsonl}")
    print(f"JSONL chunks RAG: {args.chunks_jsonl}")


if __name__ == "__main__":
    main()